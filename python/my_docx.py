# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Cm

file_name = 'test.docx'

doc = Document()

p = doc.add_paragraph('新しい段落1', style='Heading 1')



p = doc.add_paragraph('新しい段落')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True


p = doc.add_paragraph('新しい段落2', style='Heading 1')
p = doc.add_paragraph('text2')
p = doc.add_paragraph('新しい段落2-1', style='Heading 2')
p = doc.add_paragraph('text2-1')

p = doc.add_paragraph('新しい段落2-2', style='Heading 2')
p = doc.add_paragraph('text2-2')


doc.add_page_break()

bmp = 'サンプル.bmp'
doc.add_picture( bmp , width=Cm(6.3) )


doc.save(file_name)
