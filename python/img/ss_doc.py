# -*- coding: utf-8 -*-

from ss_crop import ss_crop

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


import numpy as np

import os
import glob

class ss_doc(object):
	""" Document作成クラス """
	def __init__(self, _dir='./scene'):
		""" 初期化 """
		self.doc = Document()
		
		# 処理するディレクトリ
		#date/編/シーン/xxx-com-com
		self.dir = _dir
		
	def add_header(self, _text):
		""" 見出しの挿入 """
		p = self.doc.add_paragraph(_text, style='Heading 1')
		p.style.font.bold = False
		p.style.font.size = Pt(12)
		p.style.font.color.rgb = RGBColor(0,0,0)
	
	def add_img_main(self):
		""" メイン画像の挿入 """
		# 一つ前の画像と比較する
		if self.ss_pre!=None:
			pre = np.asarray( self.ss_pre.img_y0 ).copy()
			now = np.asarray( self.ss.img_y0 ).copy()
			
			if np.array_equal(pre, now):
				return
		
		p = self.doc.add_paragraph()
		r = p.add_run()
		p.line_spacing = Pt(10)
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		r.add_picture( self.ss.get_y0(), width=Cm(9.42) )
		
	def add_img_text(self):
		""" セリフ画像の挿入 """
		p = self.doc.add_paragraph()
		r = p.add_run()
		p.line_spacing = Pt(10)
		r.add_picture( self.ss.get_y1(), width=Cm(6.3) )
	
	def add_ocr(self):
		""" OCRする """
		if True:
			# Androidではやり方わからず
			import pyocr
			tools = pyocr.get_available_tools()
			tool = tools[0]
			builder = pyocr.builders.TextBuilder(tesseract_layout=6)
			#text = tool.image_to_string(self.ss.img_y1, lang="Japanese", builder=builder)
			text = tool.image_to_string(self.ss.img_y1, lang="tessdata/jpn", builder=builder)
			
			# ある程度置換しておく
			replace_list = (
				(' ',''),
				(':',''),
				('.',''),
				('#',''),
				
				('\r',''),
				('\n',''),
				('?','？'),
				('!','！'),
				('1','１'),('2','２'),('3','３'),('4','４'),('5','５'),
				('6','６'),('7','７'),('8','８'),('9','９'),('0','０'),
				('…','・・'),
				('‥','・'),
				('-','・'),
				('・・・・・・・・・','・・・・・・'),
				('・・・・。','・・・。'),
				('(','（'),
				(')','）'),
				(']】','】'),
				('】]','】'),
				('[【','【'),
				('【[','【'),
				(']','】'),
				('[','【'),
				('|',''),
				
				('きん','さん'),
				('きさん','さん'),
				('さきん','さん'),
				('きつき','さっき'),
				('ききさま','きさま'),
				('ききま','きさま'),
				('っつ','っ'),
				('つっ','っ'),
				('つて','って'),
				('つた','った'),
				('コつ','っ'),
				('コっ','っ'),
				('つコ','っ'),
				('っコ','っ'),
				('コて','って'),
				('コた','った'),
				('つ。','っ。'),
				('やゃ','ゃ'),
				('ゃや','ゃ'),
				('ふうん','ふぅん'),
				('ふうぅん','ふぅん'),
				('ふうぅう','ふぅ'),
				('ふうぅ','ふぅ'),
				('ふぅう','ふぅ'),
				('大立夫','大丈夫'),
				('女斉','女子寮'),
				('女喜','女子寮'),
				('意庫','倉庫'),
				('むお','む'),
				('おむ','む'),
				
				('ず供','子供'),
				('ぞ供','子供'),
				
				('ポボク','ボク'),
				('ボポク','ボク'),
				('ボクタ','ボク'),
				('ボタク','ボク'),
				('ポクタ','ボク'),
				('ポタク','ボク'),
				('ポク','ボク'),
				('パカ','バカ'),
				
				('ブプリンセス','プリンセス'),
				('ブリンセス','プリンセス'),
				('アべベニュー','アベニュー'),
				('アべぺニュー','アベニュー'),
				('アぺべニュー','アベニュー'),
				('アぺぺニュー','アベニュー'),
				('センシトラル','セントラル'),
				('セシントラル','セントラル'),
				('テープル','テーブル'),
				('プラック','ブラック'),
				('曇近','最近'),
				('総画','絵画'),
				
				('境が','俺が'),
				('境は','俺は'),
				('境の','俺の'),
				('境だ','俺だ'),
				('境と','俺と'),
				('境様','俺様'),
				('届が','俺が'),
				('届は','俺は'),
				('届の','俺の'),
				('届だ','俺だ'),
				('届と','俺と'),
				('届様','俺様'),
				('恒が','俺が'),
				('恒は','俺は'),
				('恒の','俺の'),
				('恒だ','俺だ'),
				('恒と','俺と'),
				('恒様','俺様'),
				#('値が','俺が'),
				#('値は','俺は'),
				#('値の','俺の'),
				#('値だ','俺だ'),
				#('値と','俺と'),
				('値様','俺様'),
				('仙が','俺が'),
				('仙は','俺は'),
				('仙の','俺の'),
				('仙だ','俺だ'),
				('仙と','俺と'),
				('仙様','俺様'),
				('寺が','俺が'),
				('寺は','俺は'),
				('寺の','俺の'),
				('寺だ','俺だ'),
				('寺と','俺と'),
				('寺様','俺様'),
				('慎が','俺が'),
				('慎は','俺は'),
				('慎の','俺の'),
				('慎だ','俺だ'),
				('慎と','俺と'),
				('慎様','俺様'),
				('塚が','俺が'),
				('塚は','俺は'),
				('塚の','俺の'),
				('塚だ','俺だ'),
				('塚と','俺と'),
				('塚様','俺様'),
				('憧が','俺が'),
				('憧は','俺は'),
				('憧の','俺の'),
				('憧だ','俺だ'),
				('憧と','俺と'),
				('憧様','俺様'),
				
				('ちやん','ちゃん'),
				('有ゃん','ちゃん'),
				('有やん','ちゃん'),
				('台ゃん','ちゃん'),
				('台やん','ちゃん'),
				('吉ゃん','ちゃん'),
				('吉やん','ちゃん'),
				('有ゃう','ちゃう'),
				('有やう','ちゃう'),
				('台ゃう','ちゃう'),
				('台やう','ちゃう'),
				('吉ゃう','ちゃう'),
				('吉やう','ちゃう'),
				('如ゃう','ちゃう'),
				('如やう','ちゃう'),
                                
				('ババ','パパ'),

				('真紀子','真弥子'),
				('真断子','真弥子'),
				('真聞子','真弥子'),
				('真新子','真弥子'),
				('真綿子','真弥子'),
				('真絢子','真弥子'),
				('真郊子','真弥子'),
				('真線子','真弥子'),
				('真引子','真弥子'),
				('真交子','真弥子'),
				('真閑子','真弥子'),
				('真導子','真弥子'),
				('真邊子','真弥子'),
				('真部子','真弥子'),

				('真弥二','真弥子'),
				('真紀二','真弥子'),
				('真断二','真弥子'),
				('真聞二','真弥子'),
				('真新二','真弥子'),
				('真綿二','真弥子'),
				('真絢二','真弥子'),
				('真郊二','真弥子'),
				('真線二','真弥子'),
				('真引二','真弥子'),
				('真交二','真弥子'),
				('真閑二','真弥子'),
				('真導二','真弥子'),
				('真邊二','真弥子'),
				('真部二','真弥子'),

				('真弥才','真弥子'),
				('真紀才','真弥子'),
				('真断才','真弥子'),
				('真聞才','真弥子'),
				('真新才','真弥子'),
				('真綿才','真弥子'),
				('真絢才','真弥子'),
				('真郊才','真弥子'),
				('真線才','真弥子'),
				('真引才','真弥子'),
				('真交才','真弥子'),
				('真閑才','真弥子'),
				('真導才','真弥子'),
				('真邊才','真弥子'),
				('真部才','真弥子'),

				('真弥す','真弥子'),
				('真紀す','真弥子'),
				('真断す','真弥子'),
				('真聞す','真弥子'),
				('真新す','真弥子'),
				('真綿す','真弥子'),
				('真絢す','真弥子'),
				('真郊す','真弥子'),
				('真線す','真弥子'),
				('真引す','真弥子'),
				('真交す','真弥子'),
				('真閑す','真弥子'),
				('真導す','真弥子'),
				('真邊す','真弥子'),
				('真部す','真弥子'),

				('真弥地','真弥子'),
				('真紀地','真弥子'),
				('真断地','真弥子'),
				('真聞地','真弥子'),
				('真新地','真弥子'),
				('真綿地','真弥子'),
				('真絢地','真弥子'),
				('真郊地','真弥子'),
				('真線地','真弥子'),
				('真引地','真弥子'),
				('真交地','真弥子'),
				('真閑地','真弥子'),
				('真導地','真弥子'),
				('真邊地','真弥子'),
				('真部地','真弥子'),

				('真弥了','真弥子'),
				('真紀了','真弥子'),
				('真断了','真弥子'),
				('真聞了','真弥子'),
				('真新了','真弥子'),
				('真綿了','真弥子'),
				('真絢了','真弥子'),
				('真郊了','真弥子'),
				('真線了','真弥子'),
				('真引了','真弥子'),
				('真交了','真弥子'),
				('真閑了','真弥子'),
				('真導了','真弥子'),
				('真邊了','真弥子'),
				('真部了','真弥子'),

				('真弥十','真弥子'),
				('真紀十','真弥子'),
				('真断十','真弥子'),
				('真聞十','真弥子'),
				('真新十','真弥子'),
				('真綿十','真弥子'),
				('真絢十','真弥子'),
				('真郊十','真弥子'),
				('真線十','真弥子'),
				('真引十','真弥子'),
				('真交十','真弥子'),
				('真閑十','真弥子'),
				('真導十','真弥子'),
				('真邊十','真弥子'),
				('真部十','真弥子'),

				('真弥耳','真弥子'),
				('真紀耳','真弥子'),
				('真断耳','真弥子'),
				('真聞耳','真弥子'),
				('真新耳','真弥子'),
				('真綿耳','真弥子'),
				('真絢耳','真弥子'),
				('真郊耳','真弥子'),
				('真線耳','真弥子'),
				('真引耳','真弥子'),
				('真交耳','真弥子'),
				('真閑耳','真弥子'),
				('真導耳','真弥子'),
				('真邊耳','真弥子'),
				('真部耳','真弥子'),

				('真弥寺','真弥子'),
				('真紀寺','真弥子'),
				('真断寺','真弥子'),
				('真聞寺','真弥子'),
				('真新寺','真弥子'),
				('真綿寺','真弥子'),
				('真絢寺','真弥子'),
				('真郊寺','真弥子'),
				('真線寺','真弥子'),
				('真引寺','真弥子'),
				('真交寺','真弥子'),
				('真閑寺','真弥子'),
				('真導寺','真弥子'),
				('真邊寺','真弥子'),
				('真部寺','真弥子'),
				
				('真弥隔','真弥子'),
				('真紀隔','真弥子'),
				('真断隔','真弥子'),
				('真聞隔','真弥子'),
				('真新隔','真弥子'),
				('真綿隔','真弥子'),
				('真絢隔','真弥子'),
				('真郊隔','真弥子'),
				('真線隔','真弥子'),
				('真引隔','真弥子'),
				('真交隔','真弥子'),
				('真閑隔','真弥子'),
				('真導隔','真弥子'),
				('真邊隔','真弥子'),
				('真部隔','真弥子'),

				('真弥豆','真弥子'),
				('真紀豆','真弥子'),
				('真断豆','真弥子'),
				('真聞豆','真弥子'),
				('真新豆','真弥子'),
				('真綿豆','真弥子'),
				('真絢豆','真弥子'),
				('真郊豆','真弥子'),
				('真線豆','真弥子'),
				('真引豆','真弥子'),
				('真交豆','真弥子'),
				('真閑豆','真弥子'),
				('真導豆','真弥子'),
				('真邊豆','真弥子'),
				('真部豆','真弥子'),

				('真弥考','真弥子'),
				('真紀考','真弥子'),
				('真断考','真弥子'),
				('真聞考','真弥子'),
				('真新考','真弥子'),
				('真綿考','真弥子'),
				('真絢考','真弥子'),
				('真郊考','真弥子'),
				('真線考','真弥子'),
				('真引考','真弥子'),
				('真交考','真弥子'),
				('真閑考','真弥子'),
				('真導考','真弥子'),
				('真邊考','真弥子'),
				('真部考','真弥子'),

				('真弥直','真弥子'),
				('真紀直','真弥子'),
				('真断直','真弥子'),
				('真聞直','真弥子'),
				('真新直','真弥子'),
				('真綿直','真弥子'),
				('真絢直','真弥子'),
				('真郊直','真弥子'),
				('真線直','真弥子'),
				('真引直','真弥子'),
				('真交直','真弥子'),
				('真閑直','真弥子'),
				('真導直','真弥子'),
				('真邊直','真弥子'),
				('真部直','真弥子'),

				('真弥ゴ','真弥子'),
				('真紀ゴ','真弥子'),
				('真断ゴ','真弥子'),
				('真聞ゴ','真弥子'),
				('真新ゴ','真弥子'),
				('真綿ゴ','真弥子'),
				('真絢ゴ','真弥子'),
				('真郊ゴ','真弥子'),
				('真線ゴ','真弥子'),
				('真引ゴ','真弥子'),
				('真交ゴ','真弥子'),
				('真閑ゴ','真弥子'),
				('真導ゴ','真弥子'),
				('真邊ゴ','真弥子'),
				('真部ゴ','真弥子'),

				('真弥耶','真弥子'),
				('真紀耶','真弥子'),
				('真断耶','真弥子'),
				('真聞耶','真弥子'),
				('真新耶','真弥子'),
				('真綿耶','真弥子'),
				('真絢耶','真弥子'),
				('真郊耶','真弥子'),
				('真線耶','真弥子'),
				('真引耶','真弥子'),
				('真交耶','真弥子'),
				('真閑耶','真弥子'),
				('真導耶','真弥子'),
				('真邊耶','真弥子'),
				('真部耶','真弥子'),

				('真弥ず','真弥子'),
				('真紀ず','真弥子'),
				('真断ず','真弥子'),
				('真聞ず','真弥子'),
				('真新ず','真弥子'),
				('真綿ず','真弥子'),
				('真絢ず','真弥子'),
				('真郊ず','真弥子'),
				('真線ず','真弥子'),
				('真引ず','真弥子'),
				('真交ず','真弥子'),
				('真閑ず','真弥子'),
				('真導ず','真弥子'),
				('真邊ず','真弥子'),
				('真部ず','真弥子'),
				
				('真紀字','真弥子'),
				('真断字','真弥子'),
				('真聞字','真弥子'),
				('真新字','真弥子'),
				('真綿字','真弥子'),
				('真絢字','真弥子'),
				('真郊字','真弥子'),
				('真線字','真弥子'),
				('真引字','真弥子'),
				('真交字','真弥子'),
				('真閑字','真弥子'),
				('真導字','真弥子'),
				('真邊字','真弥子'),
				('真部字','真弥子'),

				('真紀孝','真弥子'),
				('真断孝','真弥子'),
				('真聞孝','真弥子'),
				('真新孝','真弥子'),
				('真綿孝','真弥子'),
				('真絢孝','真弥子'),
				('真郊孝','真弥子'),
				('真線孝','真弥子'),
				('真引孝','真弥子'),
				('真交孝','真弥子'),
				('真閑孝','真弥子'),
				('真導孝','真弥子'),
				('真邊孝','真弥子'),
				('真部孝','真弥子'),

				('真紀','真弥子'),
				('真断','真弥子'),
				('真聞','真弥子'),
				('真新','真弥子'),
				('真綿','真弥子'),
				('真絢','真弥子'),
				('真郊','真弥子'),
				('真線','真弥子'),
				('真引','真弥子'),
				('真交','真弥子'),
				('真閑','真弥子'),
				('真導','真弥子'),
				('真邊','真弥子'),
				('真部','真弥子'),

				('【真弥】','【真弥子】'),

				('法築','法条'),
				('法生','法条'),
				('法桑','法条'),

				('折堂','御堂'),
				('補堂','御堂'),
				('制堂','御堂'),
				('接堂','御堂'),

				('折党','御堂'),
				('補党','御堂'),
				('制党','御堂'),
				('接党','御堂'),

				('三階堂','二階堂'),
				('二階営','二階堂'),
				('三階党','二階堂'),
				('二階党','二階堂'),

				('三隊堂','二階堂'),
				('三隊営','二階堂'),
				('三障堂','二階堂'),
				('三障営','二階堂'),
				('三陽堂','二階堂'),
				('三陽営','二階堂'),
				('三昌堂','二階堂'),
				('三昌営','二階堂'),
				('三曲堂','二階堂'),
				('三曲営','二階堂'),
				('三降堂','二階堂'),
				('三降営','二階堂'),
				('三隆堂','二階堂'),
				('三隆営','二階堂'),
				('二隊堂','二階堂'),
				('二隊営','二階堂'),
				('二障堂','二階堂'),
				('二障営','二階堂'),
				('二陽営','二階堂'),
				('二陽堂','二階堂'),
				('二昌堂','二階堂'),
				('二昌営','二階堂'),
				('二曲堂','二階堂'),
				('二曲営','二階堂'),
				('二降堂','二階堂'),
				('二降営','二階堂'),
				('二隆堂','二階堂'),
				('二隆営','二階堂'),

				('三隊党','二階堂'),
				('三隊党','二階堂'),
				('三障党','二階堂'),
				('三障党','二階堂'),
				('三陽党','二階堂'),
				('三陽党','二階堂'),
				('三昌党','二階堂'),
				('三昌党','二階堂'),
				('三曲党','二階堂'),
				('三曲党','二階堂'),
				('三降党','二階堂'),
				('三降党','二階堂'),
				('三隆党','二階堂'),
				('三隆党','二階堂'),
				('二隊党','二階堂'),
				('二隊党','二階堂'),
				('二障党','二階堂'),
				('二障党','二階堂'),
				('二陽党','二階堂'),
				('二陽党','二階堂'),
				('二昌党','二階堂'),
				('二昌党','二階堂'),
				('二曲党','二階堂'),
				('二曲党','二階堂'),
				('二降党','二階堂'),
				('二降党','二階堂'),
				('二隆党','二階堂'),
				('二隆党','二階堂'),
				
				('小次人朗','小次郎'),
				('小次人郎','小次郎'),
				('中次朗','小次郎'),
				('中次郎','小次郎'),
				('直次朗','小次郎'),
				('直次郎','小次郎'),
				
				('小次朗','小次郎'),
				('小次帥','小次郎'),
				('小次序','小次郎'),
				('小次雇','小次郎'),
				('小次似','小次郎'),
				('小次部','小次郎'),
				('小次計','小次郎'),
				('小次分','小次郎'),
				('小次他','小次郎'),
				('小次衣','小次郎'),
				
				('小光郎','小次郎'),
				('小光朗','小次郎'),
				('小光帥','小次郎'),
				('小光序','小次郎'),
				('小光雇','小次郎'),
				('小光似','小次郎'),
				('小光部','小次郎'),
				('小光計','小次郎'),
				('小光分','小次郎'),
				('小光他','小次郎'),
				('小光衣','小次郎'),
				
				('小湊郎','小次郎'),
				('小湊朗','小次郎'),
				('小湊帥','小次郎'),
				('小湊序','小次郎'),
				('小湊雇','小次郎'),
				('小湊似','小次郎'),
				('小湊部','小次郎'),
				('小湊計','小次郎'),
				('小湊分','小次郎'),
				('小湊他','小次郎'),
				('小湊衣','小次郎'),
				
				('Eh次郎','【小次郎'),
				('Eh次朗','【小次郎'),
				('Eh次帥','【小次郎'),
				('Eh次序','【小次郎'),
				('Eh次雇','【小次郎'),
				('Eh次似','【小次郎'),
				('Eh次部','【小次郎'),
				('Eh次計','【小次郎'),
				('Eh次分','【小次郎'),
				('Eh次他','【小次郎'),
				('Eh次衣','【小次郎'),

				('Eh光郎','【小次郎'),
				('Eh光朗','【小次郎'),
				('Eh光帥','【小次郎'),
				('Eh光序','【小次郎'),
				('Eh光雇','【小次郎'),
				('Eh光似','【小次郎'),
				('Eh光部','【小次郎'),
				('Eh光計','【小次郎'),
				('Eh光分','【小次郎'),
				('Eh光他','【小次郎'),
				('Eh光衣','【小次郎'),

                                ('h次郎','小次郎'),
                                ('h次朗','小次郎'),
                                ('h次帥','小次郎'),
                                ('h次序','小次郎'),
                                ('h次雇','小次郎'),
                                ('h次似','小次郎'),
                                ('h次部','小次郎'),
                                ('h次計','小次郎'),
                                ('h次分','小次郎'),
                                ('h次他','小次郎'),
                                ('h次衣','小次郎'),

                                ('h光郎','小次郎'),
                                ('h光朗','小次郎'),
                                ('h光帥','小次郎'),
                                ('h光序','小次郎'),
                                ('h光雇','小次郎'),
                                ('h光似','小次郎'),
                                ('h光部','小次郎'),
                                ('h光計','小次郎'),
                                ('h光分','小次郎'),
                                ('h光他','小次郎'),
                                ('h光衣','小次郎'),

				('【芋】','【茜】'),
				('【再】','【茜】'),
				('【繭】','【茜】'),
				('【萌】','【茜】'),
				('【爾】','【茜】'),
				('【昔】','【茜】'),
				('【蒼】','【茜】'),
				
				('女のず】','女の子】'),
				('女のゴ】','女の子】'),
				('女のそ】','女の子】'),
				('女の】','女の子】'),
				
				('ダレンシン','グレン'),
				('ダグダレン','グレン'),
				('グダグレン','グレン'),
				('グダレジン','グレン'),
				('グレンシン','グレン'),
				('ダレッツ','グレン'),
				('グレシン','グレン'),
				('ダグレン','グレン'),
				('ダレツン','グレン'),
				('ダレンジ','グレン'),
				('ダレン','グレン'),
				
				('ウエイトヒス','ウエイトレス'),
				('うエイトレス','ウエイトレス'),
				('エイトドトヒレス','ウエイトレス'),
				('【エイトレス','【ウエイトレス'),
				
				('新生','弥生'),
				('紀生','弥生'),
				('交生','弥生'),
				('訂生','弥生'),
				('閑生','弥生'),
				('強生','弥生'),
				('絢生','弥生'),
				('引生','弥生'),
				('郊生','弥生'),
				('邊生','弥生'),
				('綿生','弥生'),
				('郊生','弥生'),
				('統生','弥生'),
				('世生','弥生'),
				('算生','弥生'),
				('尊生','弥生'),
				('聞生','弥生'),
				('線生','弥生'),
				('秋生','弥生'),
				('弾生','弥生'),

				('【乳】','【孔】'),
				('【礼】','【孔】'),
				('【攻】','【孔】'),
				('【防】','【孔】'),
				('【基】','【孔】'),
				('了乳','孔'),
				('了字','孔'),
				('了抱','孔'),
				('了凶','孔'),
				
				#('松乃','松乃'),
				('松妃','松乃'),
				('松狂','松乃'),
				('松肪','松乃'),
				('松考','松乃'),
				('松夕','松乃'),
				('松用','松乃'),

				('要乃','松乃'),
				('要妃','松乃'),
				('要狂','松乃'),
				('要肪','松乃'),
				('要考','松乃'),
				('要夕','松乃'),
				('要用','松乃'),

				('禄乃','松乃'),
				('禄妃','松乃'),
				('禄狂','松乃'),
				('禄肪','松乃'),
				('禄考','松乃'),
				('禄夕','松乃'),
				('禄用','松乃'),

				('和乃','松乃'),
				('和妃','松乃'),
				('和狂','松乃'),
				('和肪','松乃'),
				('和考','松乃'),
				('和夕','松乃'),
				('和用','松乃'),

				('栓乃','松乃'),
				('栓妃','松乃'),
				('栓狂','松乃'),
				('栓肪','松乃'),
				('栓考','松乃'),
				('栓夕','松乃'),
				('栓用','松乃'),

				('税乃','松乃'),
				('税妃','松乃'),
				('税狂','松乃'),
				('税肪','松乃'),
				('税考','松乃'),
				('税夕','松乃'),
				('税用','松乃'),

				('初乃','松乃'),
				('初妃','松乃'),
				('初狂','松乃'),
				('初肪','松乃'),
				('初考','松乃'),
				('初夕','松乃'),
				('初用','松乃'),

				('補乃','松乃'),
				('補妃','松乃'),
				('補狂','松乃'),
				('補肪','松乃'),
				('補考','松乃'),
				('補夕','松乃'),
				('補用','松乃'),

				('究乃','松乃'),
				('究妃','松乃'),
				('究狂','松乃'),
				('究肪','松乃'),
				('究考','松乃'),
				('究夕','松乃'),
				('究用','松乃'),

				('各乃','松乃'),
				('各妃','松乃'),
				('各狂','松乃'),
				('各肪','松乃'),
				('各考','松乃'),
				('各夕','松乃'),
				('各用','松乃'),

				('窟乃','松乃'),
				('窟妃','松乃'),
				('窟狂','松乃'),
				('窟肪','松乃'),
				('窟考','松乃'),
				('窟夕','松乃'),
				('窟用','松乃'),

				('柚乃','松乃'),
				('柚妃','松乃'),
				('柚狂','松乃'),
				('柚肪','松乃'),
				('柚考','松乃'),
				('柚夕','松乃'),
				('柚用','松乃'),

			)
			for (s, t) in replace_list:
				text = text.replace (s, t)
			self.add_text( text )
	
	def add_brakets(self):
		""" 段落の挿入(括弧のみ) """
		self.add_text('【】')
		
	def add_command(self, s):
		""" 段落の挿入(コマンド) """
		if len(s)>1:
			com = ''
			for c in s[1:]:
				com += c +'>'
			self.add_text(com[:-1])
				
	def add_text(self, _text):
		""" 段落の挿入 """
		p = self.doc.add_paragraph(_text, style='Body Text')
		p.paragraph_format.line_spacing = Pt(10)
		p.style.font.size = Pt(10)
	
	def run(self):
		""" 実行 """
		filename = os.path.basename( self.dir ) + '.docx'
		print(filename)
		
		# シーン毎に段落を作成する
		scene = os.path.basename( self.dir ).split('-')[1]
		self.add_header(scene)
		
		# 
		self.ss_pre = None
		
		for g in sorted(glob.glob( self.dir + '/*')):
			# シーン内のディレクトリを検索
			s = os.path.basename(g).split('-')
			
			# コマンドの書き込み
			self.add_command(s)
			
			# 各画像の検索
			for f in sorted(glob.glob(g + '/*.png')):
				self.ss = ss_crop(f)
				
				# 画像の挿入
				self.add_img_main()
				
				# セリフ画像の挿入
				self.add_img_text()
				
				# セリフ用のカッコを挿入しておく
				#self.add_brakets()
				
				# OCRする
				self.add_ocr()
				
				# 1つ前の画像を保持
				self.ss_pre = self.ss
		
		# ファイルを保存する
		self.doc.save(filename)
		
if __name__ == '__main__':
	ss_doc('./01-本部').run()
	
